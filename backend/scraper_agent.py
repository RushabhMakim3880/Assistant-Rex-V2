import os
import time
import requests
import pandas as pd
from bs4 import BeautifulSoup
from googlesearch import search
from docx import Document
from google import genai
from google.genai import types

class ScraperAgent:
    def __init__(self, project_manager):
        self.project_manager = project_manager
        # Initialize Gemini Client for data structuring (using same key as core)
        self.client = genai.Client(http_options={"api_version": "v1beta"}, api_key=os.getenv("GEMINI_API_KEY"))
        self.model = "gemini-2.0-flash-exp" # Fast model for processing

    def get_output_dir(self):
        # Create 'scraped_data' folder in current project
        path = self.project_manager.get_current_project_path() / "scraped_data"
        os.makedirs(path, exist_ok=True)
        return path

    def search_web(self, query, num_results=5):
        print(f"[Scraper] Searching for: {query}")
        try:
            # simple generator handling
            results = list(search(query, num_results=num_results))
            return results
        except Exception as e:
            print(f"[Scraper] Search failed: {e}")
            return []

    def extract_text(self, url):
        print(f"[Scraper] Extracting: {url}")
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
            response = requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Kill script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.extract()
            
            text = soup.get_text()
            # Break into lines and remove leading and trailing space on each
            lines = (line.strip() for line in text.splitlines())
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            # Drop blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text[:10000] # Limit char count for LLM
        except Exception as e:
            print(f"[Scraper] Extraction failed for {url}: {e}")
            return None

    def structure_data(self, raw_text_list, query):
        print(f"[Scraper] Structuring data with Gemini...")
        if not raw_text_list:
            return None

        combined_text = "\n\n---NEXT SOURCE---\n\n".join(raw_text_list)
        
        prompt = f"""
        You are a Data Extraction Specialist.
        User Query: "{query}"
        
        Extract relevant data from the following text sources and structure it as a JSON array of objects.
        Each object should represent a distinct item found (e.g., a company, a job, a product).
        
        If the query matches "companies", extract fields like: Name, Location, Industry, Description, Website (if found/inferred).
        If the query matches "jobs", extract: Title, Company, Location, Description.
        Generalize for other queries.
        
        SOURCES:
        {combined_text}
        
        Return ONLY valid JSON.
        """
        
        try:
            response = self.client.models.generate_content(
                model=self.model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                )
            )
            import json
            return json.loads(response.text)
        except Exception as e:
            print(f"[Scraper] Structuring failed: {e}")
            return None

    def save_data(self, data, query, format="excel"):
        if not data:
            return "No data found to save."

        # Sanitize query for filename
        filename = "".join(x for x in query if x.isalnum() or x in " -_").strip().replace(" ", "_")
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_dir = self.get_output_dir()
        
        if format.lower() == "excel":
            try:
                df = pd.DataFrame(data)
                filepath = output_dir / f"{filename}_{timestamp}.xlsx"
                df.to_excel(filepath, index=False)
                return f"Data saved to Excel: {filepath}"
            except Exception as e:
                return f"Failed to save Excel: {e}"
        
        elif format.lower() == "word":
            try:
                doc = Document()
                doc.add_heading(f"Scraped Data: {query}", 0)
                
                for item in data:
                    doc.add_heading(str(item.get('Name', item.get('Title', 'Item'))), level=1)
                    for k, v in item.items():
                        doc.add_paragraph(f"{k}: {v}")
                    doc.add_paragraph("--------------------------------------------------")
                
                filepath = output_dir / f"{filename}_{timestamp}.docx"
                doc.save(filepath)
                return f"Data saved to Word: {filepath}"
            except Exception as e:
                return f"Failed to save Word: {e}"
        else:
            return f"Unknown format: {format}"

    def run_scrape(self, query, output_format="excel"):
        # 1. Search
        urls = self.search_web(query)
        if not urls:
            return "No URLs found for this query."
        
        # 2. Extract
        raw_texts = []
        for url in urls:
            text = self.extract_text(url)
            if text:
                raw_texts.append(text)
        
        # 3. Structure
        structured_data = self.structure_data(raw_texts, query)
        
        # 4. Save
        result_msg = self.save_data(structured_data, query, output_format)
        return result_msg
